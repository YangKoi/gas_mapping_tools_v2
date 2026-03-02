import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import plotly.graph_objects as go
import matplotlib.colors as mcolors
import math
import io
import os
from PIL import Image

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from shapely.geometry import Point, LineString, Polygon
import shapely.affinity as affinity
from shapely.ops import unary_union, nearest_points
from matplotlib.path import Path

# ==========================================
# CẤU HÌNH TRANG WEB & STATE
# ==========================================
st.set_page_config(page_title="Riken Viet - True 3D Gas Mapping", layout="wide", initial_sidebar_state="expanded")

with st.sidebar:
    if os.path.exists("rkv_logo.png"):
        st.image("rkv_logo.png", use_container_width=True)
    else:
        st.markdown("### RIKEN VIET")
        
    st.header("🔄 BẢNG ĐIỀU KHIỂN")
    app_mode = st.radio("Chọn luồng công việc:", [
        "1️⃣ Thiết kế Không gian Đa lớp (3D)",
        "2️⃣ Rải nhanh trên Bản vẽ 2D (Overlay)"
    ])
    
    st.markdown("---")
    st.info("💡 **V5.1 Update:** Hỗ trợ chèn Ảnh Vật cản 2D, dựng khối 3D tự động & Khôi phục 2D Preview.")

st.title("🛡️ Riken Viet - Hệ thống Thiết kế & Dự toán Vùng phủ Khí")

# --- KHỞI TẠO DỮ LIỆU ---
if 'room_data' not in st.session_state:
    st.session_state.room_data = pd.DataFrame({"X": [0, 15, 15, 0], "Y": [0, 0, 10, 10]}) 
if 'obs_data' not in st.session_state:
    st.session_state.obs_data = pd.DataFrame([
        {"Type": "Cylinder", "Role": "Vật cản đặc (Che khí)", "X": 11.0, "Y": 8.0, "Z_base": 0.0, "Width_Radius": 1.5, "Length": 0.0, "Height": 4.0, "Angle": 0}
    ])
if 'img_obs_list' not in st.session_state:
    st.session_state.img_obs_list = [] # Lưu danh sách ảnh vật cản 2D

if 'auto_config' not in st.session_state:
    st.session_state.auto_config = pd.DataFrame([
        {"Target Gas": "CH4", "Phương thức": "Khuếch tán", "Layer": "Khí Nhẹ (Sát trần)", "Model": "SD-1", "Radius": 5.0, "Color": "cyan"},
        {"Target Gas": "H2S", "Phương thức": "Bơm hút", "Layer": "Khí Nặng (Sát sàn)", "Model": "GD-70D", "Radius": 0.0, "Color": "magenta"}
    ])
if 'leak_data' not in st.session_state:
    st.session_state.leak_data = pd.DataFrame([
        {"Tên Khu vực/Thiết bị": "Bơm hóa chất 01", "Gas": "CH4", "X": 7.5, "Y": 5.0, "Z_base": 0.5, "Bán kính rủi ro (m)": 4.0}
    ])
if 'det_data' not in st.session_state:
    st.session_state.det_data = pd.DataFrame(columns=["ID", "Model", "Gas", "Phương thức", "X", "Y", "Z", "X_hút", "Y_hút", "Z_hút", "Radius", "Color"])

# TAB 2 DATA
if 'det_data_2d' not in st.session_state:
    st.session_state.det_data_2d = pd.DataFrame(columns=["ID", "Model", "Gas", "Phương thức", "X", "Y", "X_hút", "Y_hút", "Radius", "Color"])
if 'obs_data_2d' not in st.session_state:
    st.session_state.obs_data_2d = pd.DataFrame([
        {"Type": "Box", "Role": "Vật cản đặc (Che khí)", "X": 15.0, "Y": 10.0, "Z_base": 0.0, "Width_Radius": 5.0, "Length": 5.0, "Height": 0.0, "Angle": 0}
    ])
if 'auto_config_2d' not in st.session_state:
    st.session_state.auto_config_2d = pd.DataFrame([
        {"Target Gas": "CH4", "Phương thức": "Khuếch tán", "Model": "SD-1", "Radius": 5.0, "Color": "cyan"},
        {"Target Gas": "H2S", "Phương thức": "Bơm hút", "Model": "GD-70D", "Radius": 0.0, "Color": "magenta"}
    ])
if 'leak_data_2d' not in st.session_state:
    st.session_state.leak_data_2d = pd.DataFrame(columns=["Tên Khu vực/Thiết bị", "Gas", "X", "Y", "Bán kính rủi ro (m)"])


# ==========================================
# CÁC HÀM LÕI VÀ ĐỒ HỌA
# ==========================================
def create_obstacle_polys(df_obs, img_list):
    obs_polys = []
    # 1. Vật cản hình học vẽ tay
    for _, row in df_obs.iterrows():
        if pd.isna(row['X']) or pd.isna(row['Y']) or pd.isna(row['Width_Radius']): continue
        if row.get('Role', 'Vật cản đặc (Che khí)') != 'Vật cản đặc (Che khí)': continue
            
        if row['Type'] in ['Cylinder', 'Sphere']:
            obs_polys.append(Point(row['X'], row['Y']).buffer(row['Width_Radius']))
        elif row['Type'] == 'Box':
            w = row['Width_Radius']
            l = row['Length'] if not pd.isna(row['Length']) else 0
            ang = row.get('Angle', 0) if not pd.isna(row.get('Angle', 0)) else 0
            box = Polygon([(row['X']-w/2, row['Y']-l/2), (row['X']+w/2, row['Y']-l/2),
                           (row['X']+w/2, row['Y']+l/2), (row['X']-w/2, row['Y']+l/2)])
            obs_polys.append(affinity.rotate(box, ang, origin='center'))
            
    # 2. Vật cản sinh ra từ khối Ảnh 2D
    for img_obs in img_list:
        w, l = img_obs['W'], img_obs['L']
        x, y = img_obs['X'], img_obs['Y']
        box = Polygon([(x-w/2, y-l/2), (x+w/2, y-l/2), (x+w/2, y+l/2), (x-w/2, y+l/2)])
        obs_polys.append(box)
        
    return obs_polys

def generate_2d_plot(room_poly, obs_polys, df_dets_group, df_leaks, gas_name, px, py, bg_img=None, b_w=0, b_h=0):
    minx, miny, maxx, maxy = room_poly.bounds
    res = 0.2
    xx, yy = np.meshgrid(np.arange(minx, maxx, res), np.arange(miny, maxy, res))
    pts_x, pts_y = xx.flatten(), yy.flatten()

    room_path = Path(list(room_poly.exterior.coords))
    in_room_mask = room_path.contains_points(np.c_[pts_x, pts_y])
    
    in_obs_mask = np.zeros(len(pts_x), dtype=bool)
    for obs in obs_polys:
        obs_path = Path(list(obs.exterior.coords))
        in_obs_mask |= obs_path.contains_points(np.c_[pts_x, pts_y])
    
    valid_points_mask = in_room_mask & ~in_obs_mask
    coverage_mask = np.zeros(len(pts_x), dtype=bool)

    diff_dets = df_dets_group[df_dets_group.get('Phương thức', 'Khuếch tán') == 'Khuếch tán']
    for _, det in diff_dets.iterrows():
        if pd.isna(det['X']) or pd.isna(det['Y']) or pd.isna(det['Radius']): continue
        dx, dy, dr = det['X'], det['Y'], det['Radius']
        dist_sq = (pts_x - dx)**2 + (pts_y - dy)**2
        in_radius_mask = dist_sq <= dr**2
        check_mask = valid_points_mask & in_radius_mask
        det_pt = Point(dx, dy)
        
        for i in np.where(check_mask)[0]:
            if coverage_mask[i]: continue
            line = LineString([det_pt, Point(pts_x[i], pts_y[i])])
            if not any(line.crosses(obs) for obs in obs_polys): coverage_mask[i] = True

    diem_kha_dung = np.sum(valid_points_mask)
    ty_le = (np.sum(coverage_mask) / diem_kha_dung) * 100 if diem_kha_dung > 0 else 0

    fig, ax = plt.subplots(figsize=(10, 8))
    coverage_2d = coverage_mask.reshape(xx.shape)
    
    if bg_img is not None:
        ax.imshow(bg_img, extent=[0, b_w, 0, b_h])
        ax.contourf(xx, yy, coverage_2d, levels=[0.5, 1], colors=['#00FFAA'], alpha=0.4, zorder=1)
        rx, ry = room_poly.exterior.xy
        ax.plot(rx, ry, 'k-', lw=1, zorder=2)
    else:
        ax.contourf(xx, yy, coverage_2d, levels=[0.5, 1], colors=['#A8E6CF'], alpha=0.6, zorder=1)
        rx, ry = room_poly.exterior.xy
        ax.plot(rx, ry, 'k-', lw=3, zorder=2)
        # Render ảnh vật cản vào báo cáo
        for img_obs in st.session_state.img_obs_list:
            ax.imshow(img_obs['img'], extent=[img_obs['X']-img_obs['W']/2, img_obs['X']+img_obs['W']/2, img_obs['Y']-img_obs['L']/2, img_obs['Y']+img_obs['L']/2], zorder=4, alpha=0.85)
    
    ax.plot(px, py, 's', color='red', markersize=12, markeredgecolor='black', zorder=5)
    ax.text(px + 0.3, py + 0.3, "Control Panel", color='red', fontweight='bold', zorder=6, bbox=dict(facecolor='white', alpha=0.8, edgecolor='none', pad=1))

    for obs in obs_polys:
        ox, oy = obs.exterior.xy
        ax.fill(ox, oy, color='gray', alpha=0.3, zorder=4)

    if df_leaks is not None and not df_leaks.empty:
        leaks_for_gas = df_leaks[df_leaks['Gas'] == gas_name]
        for _, leak in leaks_for_gas.iterrows():
            if pd.isna(leak['X']) or pd.isna(leak['Y']): continue
            r_risk = leak.get('Bán kính rủi ro (m)', 3.0)
            ax.add_patch(plt.Circle((leak['X'], leak['Y']), r_risk, color='red', fill=False, linestyle=':', lw=2, zorder=3))
            ax.plot(leak['X'], leak['Y'], 'x', color='red', markersize=10, markeredgewidth=2, zorder=5)
            ax.text(leak['X']+0.2, leak['Y']+0.2, f"⚠ {leak['Tên Khu vực/Thiết bị']}", fontsize=9, color='red', fontweight='bold', zorder=6)

    valid_colors = mcolors.CSS4_COLORS
    for _, det in df_dets_group.iterrows():
        if pd.isna(det['X']) or pd.isna(det['Y']): continue
        c = det['Color'].lower() if isinstance(det['Color'], str) and det['Color'].lower() in valid_colors else 'blue'
        
        if det.get('Phương thức', 'Khuếch tán') == 'Khuếch tán':
            ax.add_patch(plt.Circle((det['X'], det['Y']), det['Radius'], color=c, fill=False, linestyle='--', lw=1.5, zorder=3))
            ax.plot(det['X'], det['Y'], '^', color=c, markersize=12, markeredgecolor='black', zorder=5)
            ax.text(det['X']+0.3, det['Y']+0.3, str(det['ID']), fontsize=8, color='black', zorder=6, bbox=dict(facecolor='white', alpha=0.7, edgecolor='none', pad=1))
        else:
            hx, hy = det.get('X_hút', det['X']), det.get('Y_hút', det['Y'])
            if pd.isna(hx) or pd.isna(hy): hx, hy = det['X'], det['Y']
            ax.plot(det['X'], det['Y'], 's', color=c, markersize=10, markeredgecolor='black', zorder=5)
            ax.plot(hx, hy, 'd', color=c, markersize=8, markeredgecolor='white', zorder=5)
            ax.plot([det['X'], hx, hx], [det['Y'], det['Y'], hy], color=c, linestyle='-.', lw=2, zorder=4)

    ax.set_title(f"Bản đồ phân tích: {gas_name} | Mức an toàn (Khuếch tán): {ty_le:.1f}%", fontweight='bold')
    ax.axis('equal'); ax.grid(True, linestyle=':', alpha=0.5)
    return fig, ty_le

# ==========================================
# ĐỘNG CƠ 3D V5.1 (VOXEL & IMAGE EXTRUSION)
# ==========================================
def generate_plotly_3d_complex(room_poly, rz, obs_polys, df_obs, df_dets, df_leaks, px, py, pz, img_list):
    fig = go.Figure()
    rx, ry = room_poly.exterior.xy
    rx, ry = list(rx), list(ry)
    
    fig.add_trace(go.Scatter3d(x=rx, y=ry, z=[0]*len(rx), mode='lines', line=dict(color='white', width=4), name='Khung nhà xưởng', legendgroup='wall'))
    fig.add_trace(go.Scatter3d(x=rx, y=ry, z=[rz]*len(rx), mode='lines', line=dict(color='white', width=4), legendgroup='wall', showlegend=False))
    for x, y in zip(rx[:-1], ry[:-1]):
        fig.add_trace(go.Scatter3d(x=[x,x], y=[y,y], z=[0,rz], mode='lines', line=dict(color='white', width=2), legendgroup='wall', showlegend=False))

    fig.add_trace(go.Scatter3d(x=[px], y=[py], z=[pz], mode='markers+text', marker=dict(symbol='square', size=8, color='red'), text=["TỦ TRUNG TÂM"], textposition="top center", textfont=dict(color="red", size=12, weight="bold"), name="Tủ điều khiển"))

    def get_sphere(x0, y0, z0, r):
        u, v = np.mgrid[0:2*np.pi:20j, 0:np.pi:10j]
        x = r*np.cos(u)*np.sin(v)+x0
        y = r*np.sin(u)*np.sin(v)+y0
        z = r*np.cos(v)+z0
        z[z < 0] = 0 
        return x, y, z

    # KHỐI CẦU RÒ RỈ & ĐIỂM MÙ VOXEL
    leak_leg_added = False
    if df_leaks is not None and not df_leaks.empty:
        for _, leak in df_leaks.iterrows():
            if pd.isna(leak['X']) or pd.isna(leak['Y']): continue
            show_leak_leg = not leak_leg_added
            r_risk = leak.get('Bán kính rủi ro (m)', 3.0)
            lx, ly, lz = leak['X'], leak['Y'], leak.get('Z_base', 0.5)
            
            sx, sy, sz = get_sphere(lx, ly, lz, r_risk)
            fig.add_trace(go.Surface(x=sx, y=sy, z=sz, opacity=0.15, showscale=False, colorscale=[[0, 'red'], [1, 'red']], name="Mây Khí Độc (Rủi ro)", legendgroup='leak', showlegend=show_leak_leg))
            fig.add_trace(go.Scatter3d(x=[lx], y=[ly], z=[lz], mode='markers+text', marker=dict(symbol='x', size=6, color='red'), text=[f"⚠ {leak['Gas']}"], textposition="top center", textfont=dict(color="red", size=10), hoverinfo="text", hovertext=f"Nguồn: {leak['Tên Khu vực/Thiết bị']}", legendgroup='leak', showlegend=False))
            leak_leg_added = True
            
            diff_dets = df_dets[df_dets.get('Phương thức', 'Khuếch tán') == 'Khuếch tán']
            if not diff_dets.empty:
                res = 0.5 
                vx_grid = np.arange(lx - r_risk, lx + r_risk + res, res)
                vy_grid = np.arange(ly - r_risk, ly + r_risk + res, res)
                vz_grid = np.arange(max(0, lz - r_risk), lz + r_risk + res, res)
                
                blind_x, blind_y, blind_z = [], [], []
                
                for vx in vx_grid:
                    for vy in vy_grid:
                        for vz in vz_grid:
                            if (vx-lx)**2 + (vy-ly)**2 + (vz-lz)**2 <= r_risk**2:
                                in_obs = False
                                pt_2d = Point(vx, vy)
                                for obs_p in obs_polys:
                                    if obs_p.contains(pt_2d):
                                        in_obs = True; break
                                if in_obs: continue
                                
                                is_covered = False
                                for _, det in diff_dets.iterrows():
                                    dist_3d = math.sqrt((vx - det['X'])**2 + (vy - det['Y'])**2 + (vz - det['Z'])**2)
                                    if dist_3d <= det['Radius']:
                                        line = LineString([(det['X'], det['Y']), (vx, vy)])
                                        hit = False
                                        for obs_p in obs_polys:
                                            if line.crosses(obs_p):
                                                hit = True; break
                                        if not hit:
                                            is_covered = True; break
                                
                                if not is_covered:
                                    blind_x.append(vx); blind_y.append(vy); blind_z.append(vz)
                
                if blind_x:
                    fig.add_trace(go.Scatter3d(
                        x=blind_x, y=blind_y, z=blind_z, mode='markers',
                        marker=dict(size=3, color='#FF5500', symbol='circle', opacity=0.9),
                        name="🔥 ĐIỂM MÙ TỬ THẦN", hoverinfo='text', hovertext="Khuất bóng đầu dò!"
                    ))

    # ĐẦU DÒ
    added_gases = set()
    for _, det in df_dets.iterrows():
        if pd.isna(det['X']) or pd.isna(det['Y']) or pd.isna(det['Z']): continue
        show_leg = det['Gas'] not in added_gases
        c = det['Color']
        
        if det.get('Phương thức', 'Khuếch tán') == 'Khuếch tán':
            fig.add_trace(go.Scatter3d(x=[det['X']], y=[det['Y']], z=[det['Z']], mode='markers+text', marker=dict(symbol='circle', size=6, color='white'), text=[f"{det['ID']}"], textposition="top center", textfont=dict(color="white", size=10), name=f"Phân hệ {det['Gas']}", legendgroup=det['Gas'], showlegend=show_leg))
            sx, sy, sz = get_sphere(det['X'], det['Y'], det['Z'], det['Radius'])
            fig.add_trace(go.Surface(x=sx, y=sy, z=sz, opacity=0.1, showscale=False, colorscale=[[0, c], [1, c]], legendgroup=det['Gas'], showlegend=False, hoverinfo='skip'))
        else: # BƠM HÚT (Áp Trần)
            hx, hy, hz = det.get('X_hút', det['X']), det.get('Y_hút', det['Y']), det.get('Z_hút', det['Z'])
            if pd.isna(hx) or pd.isna(hy) or pd.isna(hz): hx, hy, hz = det['X'], det['Y'], det['Z']
            
            fig.add_trace(go.Scatter3d(x=[det['X']], y=[det['Y']], z=[det['Z']], mode='markers+text', marker=dict(symbol='square', size=6, color=c), text=[f"{det['ID']}"], textposition="top center", textfont=dict(color="white", size=10), name=f"Phân hệ {det['Gas']} (Bơm hút)", legendgroup=det['Gas'], showlegend=show_leg))
            fig.add_trace(go.Scatter3d(x=[hx], y=[hy], z=[hz], mode='markers', marker=dict(symbol='diamond', size=4, color='white', line=dict(color=c, width=2)), legendgroup=det['Gas'], showlegend=False))
            path_x = [det['X'], det['X'], hx, hx, hx]
            path_y = [det['Y'], det['Y'], det['Y'], hy, hy]
            path_z = [det['Z'], rz, rz, rz, hz]
            fig.add_trace(go.Scatter3d(x=path_x, y=path_y, z=path_z, mode='lines', line=dict(color=c, width=4, dash='dash'), legendgroup=det['Gas'], showlegend=False))

        added_gases.add(det['Gas'])

    # VẬT CẢN VẼ TAY CƠ BẢN
    obs_leg_added = False
    for i, (_, obs) in enumerate(df_obs.iterrows()):
        if pd.isna(obs['X']) or pd.isna(obs['Y']) or pd.isna(obs['Width_Radius']): continue
        show_obs_leg = not obs_leg_added
        z_base = obs.get('Z_base', 0.0)
        h = obs.get('Height', 4.0)
        solid_color = '#555555' if obs.get('Role') == 'Vật cản đặc (Che khí)' else 'lightblue'
        
        if obs['Type'] == 'Cylinder':
            z_grid, theta = np.mgrid[z_base:(z_base+h):2j, 0:2*np.pi:20j]
            x_cyl = obs['Width_Radius']*np.cos(theta)+obs['X']
            y_cyl = obs['Width_Radius']*np.sin(theta)+obs['Y']
            fig.add_trace(go.Surface(x=x_cyl, y=y_cyl, z=z_grid, opacity=1.0, showscale=False, colorscale=[[0, solid_color], [1, solid_color]], name="Cấu trúc Vẽ tay", legendgroup='obs', showlegend=show_obs_leg))
            obs_leg_added = True
        elif obs['Type'] == 'Box':
            w = obs['Width_Radius']
            l = obs['Length'] if not pd.isna(obs['Length']) else 0
            ang = obs.get('Angle', 0) if not pd.isna(obs.get('Angle', 0)) else 0
            box_poly = Polygon([(obs['X']-w/2, obs['Y']-l/2), (obs['X']+w/2, obs['Y']-l/2),
                           (obs['X']+w/2, obs['Y']+l/2), (obs['X']-w/2, obs['Y']+l/2)])
            bx, by = affinity.rotate(box_poly, ang, origin='center').exterior.xy
            bx, by = list(bx)[:-1], list(by)[:-1] 
            x_box, y_box = bx * 2, by * 2
            z_box = [z_base]*4 + [z_base + h]*4
            ii, jj, kk = [7,0,0,0,4,4,6,6,4,0,3,2], [3,4,1,2,5,6,5,2,0,1,6,3], [0,7,2,3,6,7,1,1,5,5,7,6]
            fig.add_trace(go.Mesh3d(x=x_box, y=y_box, z=z_box, i=ii, j=jj, k=kk, color=solid_color, opacity=1.0, name="Cấu trúc Vẽ tay", legendgroup='obs', showlegend=show_obs_leg))
            obs_leg_added = True

    # DỰNG HÌNH VẬT CẢN TỪ ẢNH 2D (EXTRUSION)
    for img_obs in img_list:
        w, l, h = img_obs['W'], img_obs['L'], img_obs['H']
        xc, yc, zb = img_obs['X'], img_obs['Y'], img_obs['Z_base']
        
        xb = [xc-w/2, xc+w/2, xc+w/2, xc-w/2, xc-w/2, xc+w/2, xc+w/2, xc-w/2]
        yb = [yc-l/2, yc-l/2, yc+l/2, yc+l/2, yc-l/2, yc-l/2, yc+l/2, yc+l/2]
        zb_arr = [zb, zb, zb, zb, zb+h, zb+h, zb+h, zb+h]
        ii, jj, kk = [7,0,0,0,4,4,6,6,4,0,3,2], [3,4,1,2,5,6,5,2,0,1,6,3], [0,7,2,3,6,7,1,1,5,5,7,6]
        
        fig.add_trace(go.Mesh3d(x=xb, y=yb, z=zb_arr, i=ii, j=jj, k=kk, color='#8899A6', opacity=0.9, name="Thiết bị (Dựng từ Ảnh)", flatshading=True))

    minx, miny, maxx, maxy = room_poly.bounds
    fig.update_layout(
        scene=dict(xaxis=dict(range=[minx, maxx], title='X', backgroundcolor="rgb(30,30,30)", gridcolor="gray"),
                   yaxis=dict(range=[miny, maxy], title='Y', backgroundcolor="rgb(30,30,30)", gridcolor="gray"),
                   zaxis=dict(range=[0, max(rz, 5)], title='Z', backgroundcolor="rgb(30,30,30)", gridcolor="gray"),
                   aspectmode='data'),
        paper_bgcolor="rgb(15,15,15)", plot_bgcolor="rgb(15,15,15)", margin=dict(l=0, r=0, b=0, t=30),
        showlegend=True, legend=dict(yanchor="top", y=0.99, xanchor="left", x=0.01, bgcolor="rgba(30, 30, 30, 0.7)", bordercolor="gray", borderwidth=1, font=dict(color="white"))
    )
    return fig

# (Các hàm build_bom_df và generate_full_word_report giữ nguyên)
def build_bom_df(det_df, p_x, p_y, r_z, p_z, waste):
    bom_items = [{"STT": 1, "Hạng mục thiết bị": "Tủ điều khiển trung tâm đo khí", "Đơn vị": "Bộ", "Khối lượng": 1}]
    stt = 2
    total_cable_length = 0
    total_tube_length = 0
    valid_dets = det_df.dropna(subset=['X', 'Y']) 
    
    if not valid_dets.empty and 'Model' in valid_dets.columns:
        counts = valid_dets['Model'].value_counts()
        for model, qty in counts.items():
            bom_items.append({"STT": stt, "Hạng mục thiết bị": f"Đầu dò đo khí rò rỉ - Model: {model}", "Đơn vị": "Bộ", "Khối lượng": qty})
            stt += 1
        for _, d in valid_dets.iterrows():
            cable_up = abs(r_z - p_z) if 'Z' in valid_dets.columns and not pd.isna(d.get('Z')) else 3.0
            cable_horizontal = abs(d['X'] - p_x) + abs(d['Y'] - p_y)
            cable_down = abs(r_z - d['Z']) if 'Z' in valid_dets.columns and not pd.isna(d.get('Z')) else 1.5 
            total_cable_length += (cable_up + cable_horizontal + cable_down)
            
            if d.get('Phương thức', 'Khuếch tán') == 'Bơm hút':
                hx, hy, hz = d.get('X_hút', d['X']), d.get('Y_hút', d['Y']), d.get('Z_hút', d['Z'])
                if pd.notna(hx) and pd.notna(hy) and pd.notna(hz):
                    tube_up = abs(r_z - d['Z'])          
                    tube_horiz = abs(hx - d['X']) + abs(hy - d['Y']) 
                    tube_down = abs(r_z - hz)            
                    total_tube_length += (tube_up + tube_horiz + tube_down)

        total_cable_length = math.ceil(total_cable_length * (1 + waste / 100))
        total_tube_length = math.ceil(total_tube_length * (1 + waste / 100))
        
    bom_items.append({"STT": stt, "Hạng mục thiết bị": "Cáp tín hiệu chống nhiễu chuyên dụng", "Đơn vị": "Mét", "Khối lượng": total_cable_length})
    stt+=1
    if total_tube_length > 0:
        bom_items.append({"STT": stt, "Hạng mục thiết bị": "Ống dẫn khí lấy mẫu chuyên dụng (Teflon/PU)", "Đơn vị": "Mét", "Khối lượng": total_tube_length})
        stt+=1
        
    bom_items.append({"STT": stt, "Hạng mục thiết bị": "Chuông đèn cảnh báo (Siren/Light)", "Đơn vị": "Bộ", "Khối lượng": len(valid_dets) if not valid_dets.empty else 1})
    return pd.DataFrame(bom_items)

def generate_full_word_report(figs_dict, img_3d_bytes, bom_df, p_name, c_name, author, r_date, r_num, is_3d=True):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)

    section = doc.sections[0]
    header = section.header
    h_para = header.paragraphs[0]
    h_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists("header_logo.png"): h_para.add_run().add_picture("header_logo.png", width=Inches(6.5))
    else: h_para.add_run("CÔNG TY TNHH CÔNG NGHỆ THIẾT BỊ DÒ KHÍ RIKEN VIET").bold = True
    
    footer = section.footer
    f_para = footer.paragraphs[0]
    f_para.text = "Bản quyền thuộc về Riken Việt, không sao chép và sử dụng sai mục đích."
    f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in f_para.runs: run.font.name = 'Arial'; run.font.size = Pt(9); run.font.italic = True; run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()
    title = doc.add_heading('BÁO CÁO THIẾT KẾ VÀ DỰ TOÁN HỆ THỐNG ĐO KHÍ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs: run.font.name = 'Arial'
    
    p_num = doc.add_paragraph(f"Số/ No.: {r_num}")
    p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p_num.runs: run.font.name = 'Arial'; run.font.italic = True
        
    doc.add_paragraph(f"Tên Dự án: {p_name}").runs[0].font.bold = True
    doc.add_paragraph(f"Đơn vị / Khách hàng: {c_name}").runs[0].font.bold = True
    doc.add_paragraph(f"Người lập báo cáo: {author}")
    doc.add_paragraph(f"Ngày xuất báo cáo: {r_date.strftime('%d/%m/%Y')}")
    doc.add_paragraph("_" * 60)
    
    section_num = 1
    if is_3d and img_3d_bytes:
        h1 = doc.add_heading(f'{section_num}. PHÂN BỔ KHÔNG GIAN TỔNG THỂ (MÔ PHỎNG 3D)', level=1)
        for run in h1.runs: run.font.name = 'Arial'
        doc.add_paragraph("Sơ đồ 3D: Vùng rủi ro (Khối cầu đỏ), Đầu dò (vòng cầu/đường ống) và Phân tích điểm mù Voxel (Chấm cam).")
        p = doc.add_paragraph()
        p.add_run().add_picture(img_3d_bytes, width=Inches(6.0))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        section_num += 1
        
    h2 = doc.add_heading(f'{section_num}. PHÂN TÍCH ĐIỂM MÙ (MẶT BẰNG 2D)', level=1)
    for run in h2.runs: run.font.name = 'Arial'
    for gas_name, fig_info in figs_dict.items():
        p_title = doc.add_paragraph()
        p_title.add_run(f"Bản đồ phân hệ: {gas_name} ").bold = True
        p_title.add_run(f"(Mức an toàn vùng khuếch tán: {fig_info['coverage']:.1f}%)").italic = True
        img_stream = io.BytesIO()
        fig_info['fig'].savefig(img_stream, format='png', bbox_inches='tight', dpi=150)
        img_stream.seek(0)
        p_img = doc.add_paragraph()
        p_img.add_run().add_picture(img_stream, width=Inches(6.0))
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    section_num += 1
        
    h3 = doc.add_heading(f'{section_num}. BẢNG BÓC TÁCH KHỐI LƯỢNG VẬT TƯ (BOM)', level=1)
    for run in h3.runs: run.font.name = 'Arial'
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    headers = ['STT', 'Hạng mục Thiết bị & Vật tư', 'Đơn vị', 'Khối lượng']
    for i in range(4):
        hdr[i].text = headers[i]
        for run in hdr[i].paragraphs[0].runs: run.font.name = 'Arial'; run.font.bold = True
    for _, row in bom_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text, row_cells[1].text, row_cells[2].text, row_cells[3].text = str(row['STT']), str(row['Hạng mục thiết bị']), str(row['Đơn vị']), f"{int(row['Khối lượng']):,}"
        for cell in row_cells:
            for p in cell.paragraphs:
                for run in p.runs: run.font.name = 'Arial'
        
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream

# ==========================================
# LUỒNG CHÍNH 1: MÔ PHỎNG 3D CHUYÊN SÂU
# ==========================================
if app_mode == "1️⃣ Thiết kế Không gian Đa lớp (3D)":
    st.markdown("## 🧊 Chế độ Thiết kế Không gian 3D Chuyên sâu")
    
    col_input1, col_input2 = st.columns([1.2, 1.1])
    
    with col_input1:
        st.header("1. Không gian & Tủ Trung Tâm")
        room_z = st.number_input("Chiều cao trần nhà xưởng (Z) - mét", min_value=1.0, value=5.0, key="z_3d")

        st.subheader("🎛️ Vị trí Tủ Điều Khiển & Hao hụt")
        col_p1, col_p2, col_p3 = st.columns(3)
        panel_x = col_p1.number_input("Tọa độ X (Tủ)", value=0.0, key="px_3d")
        panel_y = col_p2.number_input("Tọa độ Y (Tủ)", value=0.0, key="py_3d")
        panel_z = col_p3.number_input("Cao độ Z (Tủ)", value=1.5, key="pz_3d")
        wastage_percent = st.number_input("Hệ số hao hụt cáp thi công (%)", min_value=0, value=20, step=5, key="wast_3d")

        st.subheader("📐 Định hình Không gian")
        col_t1, col_t2, col_t3 = st.columns(3)
        with col_t1:
            if st.button("🟩 Chữ Nhật", use_container_width=True, key="btn_rect"): st.session_state.room_data = pd.DataFrame({"X": [0, 15, 15, 0], "Y": [0, 0, 10, 10]}); st.rerun()
        with col_t2:
            if st.button("▛ Chữ L", use_container_width=True, key="btn_l"): st.session_state.room_data = pd.DataFrame({"X": [0, 15, 15, 6, 6, 0], "Y": [0, 0, 6, 6, 12, 12]}); st.rerun()
        with col_t3:
            if st.button("⨆ Chữ U", use_container_width=True, key="btn_u"): st.session_state.room_data = pd.DataFrame({"X": [0, 20, 20, 15, 15, 5, 5, 0], "Y": [0, 0, 15, 15, 5, 5, 15, 15]}); st.rerun()

        with st.form("form_room_3d"):
            edited_room = st.data_editor(st.session_state.room_data, num_rows="dynamic", use_container_width=True, key="ed_room_3d")
            if st.form_submit_button("🔄 Vẽ lại Phòng"):
                st.session_state.room_data = edited_room
                st.rerun()
        
        # --- BẢN XEM TRƯỚC (LIVE 2D PREVIEW) ĐÃ ĐƯỢC KHÔI PHỤC ---
        st.markdown("### 👁️ Bản xem trước (Live 2D Preview)")
        if len(st.session_state.room_data) >= 3:
            room_poly = Polygon(list(zip(st.session_state.room_data['X'], st.session_state.room_data['Y'])))
            fig_grid, ax_grid = plt.subplots(figsize=(6, 5))
            x_ext, y_ext = room_poly.exterior.xy
            ax_grid.plot(x_ext, y_ext, color='#333333', linewidth=2)
            ax_grid.fill(x_ext, y_ext, alpha=0.1, color='blue')
            
            for idx, row in st.session_state.room_data.iterrows():
                if pd.isna(row['X']) or pd.isna(row['Y']): continue
                ax_grid.plot(row['X'], row['Y'], 'ro', markersize=6, zorder=10)
                ax_grid.text(row['X'] + 0.3, row['Y'] + 0.3, f"P{idx}", color='red', fontweight='bold', fontsize=11, zorder=11)
            ax_grid.plot(panel_x, panel_y, 's', color='red', markersize=12, markeredgecolor='black', zorder=12)
            
            # 1. Vẽ các ảnh 2D được upload
            for img_obs in st.session_state.img_obs_list:
                ax_grid.imshow(img_obs['img'], extent=[img_obs['X']-img_obs['W']/2, img_obs['X']+img_obs['W']/2, 
                                                       img_obs['Y']-img_obs['L']/2, img_obs['Y']+img_obs['L']/2], 
                               zorder=4, alpha=0.85)
            
            # 2. Vẽ hình học cơ bản
            obs_polys_draft = create_obstacle_polys(st.session_state.obs_data, [])
            for obs in obs_polys_draft:
                ox, oy = obs.exterior.xy
                ax_grid.fill(ox, oy, color='gray', alpha=0.4, hatch='//', zorder=5)
            
            for _, leak in st.session_state.leak_data.iterrows():
                if pd.isna(leak['X']) or pd.isna(leak['Y']): continue
                ax_grid.plot(leak['X'], leak['Y'], 'rx', markersize=10, zorder=15)
                ax_grid.add_patch(plt.Circle((leak['X'], leak['Y']), leak.get('Bán kính rủi ro (m)', 3.0), color='red', fill=False, linestyle=':', zorder=14))
                
            for _, det in st.session_state.det_data.iterrows():
                if pd.isna(det['X']) or pd.isna(det['Y']): continue
                c = det['Color'].lower() if isinstance(det['Color'], str) and det['Color'].lower() in mcolors.CSS4_COLORS else 'blue'
                if det.get('Phương thức') == 'Bơm hút':
                    hx, hy = det.get('X_hút', det['X']), det.get('Y_hút', det['Y'])
                    ax_grid.plot(det['X'], det['Y'], 's', color=c, markersize=8, markeredgecolor='black', zorder=16)
                    ax_grid.plot(hx, hy, 'd', color=c, markersize=6, zorder=16)
                    ax_grid.plot([det['X'], hx, hx], [det['Y'], det['Y'], hy], color=c, linestyle='-.', zorder=15)
                else:
                    ax_grid.plot(det['X'], det['Y'], '^', color=c, markersize=8, markeredgecolor='black', zorder=16)

            ax_grid.set_aspect('equal'); ax_grid.grid(True, linestyle='--', alpha=0.5)
            st.pyplot(fig_grid)
        else:
            st.error("Phòng cần ít nhất 3 góc!")
            room_poly = None

    with col_input2:
        st.header("2. Phân bổ Mục tiêu & Thiết bị")
        
        with st.form("form_auto_3d"):
            st.write("🚨 **Khai báo Nguồn Rủi ro (Smart Target)**")
            edited_leak_3d = st.data_editor(st.session_state.leak_data, num_rows="dynamic", use_container_width=True)

            st.write("⚙️ **Cấu hình Phân hệ Khí (Rải tự động)**")
            edited_auto_config = st.data_editor(st.session_state.auto_config, num_rows="dynamic", use_container_width=True,
                column_config={
                    "Phương thức": st.column_config.SelectboxColumn("Phương thức", options=["Khuếch tán", "Bơm hút"]),
                    "Layer": st.column_config.SelectboxColumn("Mặt phẳng", options=["Khí Nhẹ (Sát trần)", "Khí Trung bình (Vùng thở)", "Khí Nặng (Sát sàn)"]), 
                    "Color": st.column_config.SelectboxColumn("Màu", options=["cyan", "magenta", "yellow", "lime", "red"])})

            if st.form_submit_button("🚀 LƯU & TỰ ĐỘNG RẢI SMART TARGET", type="primary"):
                st.session_state.leak_data = edited_leak_3d
                st.session_state.auto_config = edited_auto_config
                
                obs_polys = create_obstacle_polys(st.session_state.obs_data, st.session_state.img_obs_list)
                obs_union = unary_union(obs_polys) if obs_polys else Polygon()
                
                if room_poly is not None:
                    new_dets = []
                    for _, row_cfg in edited_auto_config.iterrows():
                        target_gas = row_cfg["Target Gas"]
                        phuong_thuc = row_cfg.get("Phương thức", "Khuếch tán")
                        
                        if phuong_thuc == "Khuếch tán":
                            if "Nhẹ" in row_cfg["Layer"]: z_val = max(room_z - 0.5, 0.5)
                            elif "Nặng" in row_cfg["Layer"]: z_val = 0.5
                            else: z_val = 1.5 
                            spacing = row_cfg["Radius"] * 1.5 
                            
                            target_leaks = edited_leak_3d[edited_leak_3d['Gas'] == target_gas].dropna(subset=['X', 'Y'])
                            if not target_leaks.empty:
                                risk_areas = [Point(r['X'], r['Y']).buffer(r.get('Bán kính rủi ro (m)', 3.0)) for _, r in target_leaks.iterrows()]
                                valid_area = room_poly.intersection(unary_union(risk_areas))
                            else:
                                valid_area = room_poly 
                            
                            valid_area = valid_area.difference(obs_union)
                            
                            if not valid_area.is_empty:
                                minx, miny, maxx, maxy = valid_area.bounds
                                grid_res = 0.5 
                                xs = np.arange(minx, maxx + grid_res, grid_res)
                                ys = np.arange(miny, maxy + grid_res, grid_res)
                                
                                placed_pts = []
                                count = 1
                                for x in xs:
                                    for y in ys:
                                        pt = Point(x, y)
                                        if valid_area.contains(pt):
                                            too_close = False
                                            for px, py in placed_pts:
                                                if math.hypot(x - px, y - py) < spacing:
                                                    too_close = True; break
                                            if not too_close:
                                                placed_pts.append((x, y))
                                                new_dets.append({"ID": f"{row_cfg['Model']} ({count:02d})", "Model": row_cfg['Model'], "Gas": target_gas, "Phương thức": "Khuếch tán", "X": round(x, 1), "Y": round(y, 1), "Z": z_val, "X_hút": None, "Y_hút": None, "Z_hút": None, "Radius": row_cfg["Radius"], "Color": row_cfg["Color"]})
                                                count += 1
                        else: 
                            target_leaks = edited_leak_3d[edited_leak_3d['Gas'] == target_gas].dropna(subset=['X', 'Y'])
                            count = 1
                            for _, leak in target_leaks.iterrows():
                                l_pt = Point(leak['X'], leak['Y'])
                                lz = leak.get('Z_base', 0.5)
                                
                                if not obs_union.is_empty and obs_union.contains(l_pt):
                                    snap_pt = nearest_points(obs_union.boundary, l_pt)[0]
                                    sx, sy = snap_pt.x, snap_pt.y
                                else:
                                    sx, sy = leak['X'], leak['Y']
                                    
                                if not room_poly.boundary.is_empty:
                                    wall_pt = nearest_points(room_poly.boundary, Point(sx, sy))[0]
                                    mx, my = wall_pt.x, wall_pt.y
                                else:
                                    mx, my = sx, sy 
                                
                                new_dets.append({"ID": f"{row_cfg['Model']} ({count:02d})", "Model": row_cfg['Model'], "Gas": target_gas, "Phương thức": "Bơm hút", "X": round(mx, 1), "Y": round(my, 1), "Z": 1.5, "X_hút": round(sx, 1), "Y_hút": round(sy, 1), "Z_hút": lz, "Radius": row_cfg["Radius"], "Color": row_cfg["Color"]})
                                count += 1

                    st.session_state.det_data = pd.DataFrame(new_dets)
                    st.rerun()

        # --- HỆ THỐNG VẬT CẢN LAI (HYBRID OBSTACLES) ---
        st.write("🚧 **Hệ thống Lắp ghép Không gian 3D**")
        tab_basic, tab_img = st.tabs(["📐 Vẽ Khối Cơ bản", "🖼️ Nhập Ảnh Mặt Bằng 2D"])
        
        with tab_basic:
            with st.form("form_manual_3d"):
                edited_obs = st.data_editor(st.session_state.obs_data, num_rows="dynamic", use_container_width=True, 
                    column_config={
                        "Type": st.column_config.SelectboxColumn("Hình khối", options=["Cylinder", "Box", "Sphere"]),
                        "Role": st.column_config.SelectboxColumn("Thuộc tính", options=["Vật cản đặc (Che khí)", "Trang trí 3D (Xuyên thấu)"])
                    })
                st.write("📋 **Tọa độ Thiết bị:**")
                edited_dets = st.data_editor(st.session_state.det_data, num_rows="dynamic", use_container_width=True,
                    column_config={"Phương thức": st.column_config.SelectboxColumn("Phương thức", options=["Khuếch tán", "Bơm hút"])})
                
                if st.form_submit_button("🔄 XÁC NHẬN & CẬP NHẬT BẢN VẼ", type="secondary"):
                    st.session_state.obs_data = edited_obs
                    st.session_state.det_data = edited_dets
                    st.rerun()

        with tab_img:
            with st.expander("➕ Thêm mới Vật cản từ Ảnh", expanded=True):
                up_file = st.file_uploader("Tải lên ảnh cắt CAD (PNG/JPG)", type=['png', 'jpg', 'jpeg'])
                col_i1, col_i2, col_i3 = st.columns(3)
                io_x = col_i1.number_input("Tọa độ X", value=5.0)
                io_y = col_i2.number_input("Tọa độ Y", value=5.0)
                io_z = col_i3.number_input("Cao độ Đáy (Z)", value=0.0)
                
                col_i4, col_i5, col_i6 = st.columns(3)
                io_w = col_i4.number_input("Chiều Rộng X (m)", value=3.0)
                io_l = col_i5.number_input("Chiều Dài Y (m)", value=4.0)
                io_h = col_i6.number_input("Chiều Cao 3D (m)", value=3.0)
                
                if st.button("Lưu Hình ảnh vào Lưới", type="primary", use_container_width=True):
                    if up_file:
                        img = Image.open(up_file)
                        st.session_state.img_obs_list.append({
                            'id': up_file.name, 'img': img, 'X': io_x, 'Y': io_y,
                            'W': io_w, 'L': io_l, 'Z_base': io_z, 'H': io_h
                        })
                        st.rerun()
            
            if st.session_state.img_obs_list:
                st.write(f"Đang có **{len(st.session_state.img_obs_list)}** ảnh vật cản.")
                if st.button("🗑️ Xóa tất cả Ảnh", type="secondary"):
                    st.session_state.img_obs_list = []
                    st.rerun()

    st.markdown("---")
    st.header("3. 📊 Phân tích & Xuất Báo cáo")
    col_info1, col_info2, col_info3 = st.columns([1, 1, 1])
    with col_info1: project_name = st.text_input("Tên Dự án", value="Thiết kế Giám sát Rò rỉ", key="pn_1")
    with col_info2: client_name = st.text_input("Khách hàng", value="Nhà máy ABC", key="cn_1")
    with col_info3: report_number = st.text_input("Số Báo cáo", value="RKV_TE_001/BC", key="rn_1")
    
    author_name = st.text_input("Người lập báo cáo", value="Cao Minh Lợi - Giám đốc Kỹ thuật", key="au_1")
    report_date = st.date_input("Ngày lập báo cáo", key="rd_1")

    if st.button("🚀 CHẠY MÔ PHỎNG & TẠO FILE BÁO CÁO (3D)", type='primary', use_container_width=True):
        room_poly = Polygon(list(zip(st.session_state.room_data['X'], st.session_state.room_data['Y'])))
        if room_poly is None or st.session_state.det_data.dropna(subset=['X', 'Y']).empty:
            st.error("Lỗi: Cần vẽ phòng và rải đầu dò trước khi chạy!")
        else:
            with st.spinner('Đang quét lưới Voxel đa chiều và dựng hình khối 3D...'):
                obs_polys = create_obstacle_polys(st.session_state.obs_data, st.session_state.img_obs_list)
                
                fig_3d = generate_plotly_3d_complex(room_poly, room_z, obs_polys, st.session_state.obs_data, st.session_state.det_data, st.session_state.leak_data, panel_x, panel_y, panel_z, st.session_state.img_obs_list)
                st.plotly_chart(fig_3d, use_container_width=True)
                
                img_3d_bytes = io.BytesIO()
                try: fig_3d.write_image(img_3d_bytes, format='png', width=800, height=500); img_3d_bytes.seek(0)
                except: img_3d_bytes = None

                gas_groups = st.session_state.det_data.dropna(subset=['Gas'])['Gas'].unique()
                ui_tabs = st.tabs([f"Bản đồ: {g}" for g in gas_groups])
                figs_dict = {} 
                
                for i, gas_name in enumerate(gas_groups):
                    with ui_tabs[i]:
                        f, c = generate_2d_plot(room_poly, obs_polys, st.session_state.det_data[st.session_state.det_data['Gas']==gas_name], st.session_state.leak_data, gas_name, panel_x, panel_y)
                        st.pyplot(f) 
                        st.success(f"✅ Diện tích an toàn (Khuếch tán): {c:.1f}% | Bơm hút: Giám sát cục bộ tại nguồn.")
                        figs_dict[gas_name] = {'fig': f, 'coverage': c}

                bom_df = build_bom_df(st.session_state.det_data, panel_x, panel_y, room_z, panel_z, wastage_percent)
                st.write("📋 **Bảng Bóc tách Khối lượng Vật tư & Đường Ống (BOM):**")
                st.dataframe(bom_df) 
                
                word_stream = generate_full_word_report(figs_dict, img_3d_bytes, bom_df, project_name, client_name, author_name, report_date, report_number, is_3d=True)
                st.download_button("📥 TẢI BÁO CÁO WORD (3D)", word_stream, f"BaoCao_3D_V5_{client_name}.docx", type="primary")

# ==========================================
# LUỒNG CHÍNH 2: RẢI NHANH TRÊN BẢN VẼ 2D
# ==========================================
elif app_mode == "2️⃣ Rải nhanh trên Bản vẽ 2D (Overlay)":
    st.markdown("## 🖼️ Chế độ Rải nhanh trên Bản vẽ 2D (Overlay)")
    st.info("Tính năng đang bảo trì, vui lòng quay lại chế độ 3D để trải nghiệm!")
# (Đã ẩn phần Overlay để rút ngắn file. Hãy copy phần cũ vào nếu bạn cần dùng song song)

# ==========================================
# FOOTER BẢN QUYỀN
# ==========================================
st.markdown("""
    <hr style="border: 0; height: 1px; background-image: linear-gradient(to right, rgba(255, 255, 255, 0), rgba(255, 255, 255, 0.2), rgba(255, 255, 255, 0)); margin-top: 50px;">
    <div style="text-align: center; color: #888888; font-size: 14px; padding-bottom: 20px;">
        &copy; 2026 All Rights Reserved.<br>
        Designed and programmed by <b>trggiang</b>.
    </div>
""", unsafe_allow_html=True)
